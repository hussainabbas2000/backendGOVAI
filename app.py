import os
import io
import time
import tempfile
import requests
from flask import Flask, request, jsonify, Blueprint
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
import json
import re
from flask_cors import CORS
from quote import quote_bp, init_db
from notifications import notifications_bp, init_notifications_db
from compliance import compliance_bp, init_compliance_db
from pdf_generator import pdf_bp, init_pdf_db
from email_webhook import webhook_bp, init_webhook_db
from email_poller import init_poller_db
from background_jobs import init_background_jobs, get_scheduler_status, start_negotiation_for_session
from suggestions import suggestions_bp

load_dotenv()
app = Flask(__name__)
# Configure database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///sam_gov_negotiations.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize database
db = SQLAlchemy(app)

# Initialize the quote blueprint with the database
init_db(db)

# Initialize the notifications blueprint with the database
init_notifications_db(db)

# Initialize the compliance blueprint with the database
init_compliance_db(db)

# Initialize the PDF blueprint with the database
init_pdf_db(db)

# Register the blueprints
app.register_blueprint(quote_bp)
app.register_blueprint(notifications_bp)
app.register_blueprint(compliance_bp)
app.register_blueprint(pdf_bp)
app.register_blueprint(webhook_bp)
app.register_blueprint(suggestions_bp)

from quote import NegotiationSession, Supplier, Message
from notifications import Notification

# Initialize webhook after models are available
init_webhook_db(db, Supplier, Message, NegotiationSession)

# Initialize Gmail inbox poller
init_poller_db(db, Supplier, Message, NegotiationSession)


# Background job control routes
@app.route('/api/jobs/status', methods=['GET'])
def job_status():
    """Get background job scheduler status"""
    status = get_scheduler_status()
    return jsonify(status)


@app.route('/api/jobs/process-session/<int:session_id>', methods=['POST'])
def process_session(session_id):
    """Manually trigger processing for a negotiation session"""
    result = start_negotiation_for_session(session_id)
    return jsonify(result)


CORS(app, resources={
    r"/analyze-solicitations": {
        "origins": [
            "http://localhost:9002",                     # local dev
            r"^http://127\.0\.0\.1(:[0-9]+)?$",
            "https://sam-gov-liard.vercel.app"              # deployed frontend
        ]
    },
    r"/api/*": {  # This covers all quote and notifications blueprint routes
        "origins": [
            "http://localhost:3000",                      # Next.js local dev
            "http://localhost:9002",                      # your existing local dev
            "https://sam-gov-liard.vercel.app",            # deployed frontend
            "*"  # You can restrict this in production
        ]
    },
    r"/message-chat": {
        "origins": [
            "http://localhost:3000",
            "http://localhost:9002",
            "https://sam-gov-liard.vercel.app",
            "*"
        ]
    },
    r"/webhook/*": {  # Email webhook endpoints
        "origins": "*"
    }
})
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Your detailed instruction prompt
SYSTEM_PROMPT = {
    "type": "input_text",
    "text": """
    You are now acting as my full-spectrum government contracting and sourcing expert. You specialize in working across all U.S. government agencies — including civilian, defense, and intelligence sectors — and you're proficient in interpreting solicitations from federal, state, and local governments.

Your job is to assist me through the entire sourcing and procurement process for government contracts. I will provide you with solicitation docs, and you will:

🔍 1. Read and Interpret the Solicitations
Analyze any solicitation documents (RFI, RFQ, RFP, IDIQ, BPA, etc.) and provide a plain-English summary of the requirement.

Identify the agency issuing it and explain the mission or goal of the procurement.

Highlight any critical requirements, special instructions, and mandatory compliance points.

🧠 2. Extract All Key Contract Data
From the solicitation, extract and organize the following:

Solicitation Number

NAICS Code(s)

PSC/FSC Code (if listed)

Set-aside Type (e.g., Small Business, 8(a), SDVOSB)

Contract Type (e.g., FFP, IDIQ, GSA Schedule)

Response Deadline and Submission Method

Evaluation Criteria (technical, price, past performance, etc.)

Product or Service Description (this is must)

Quantity, Units, and Delivery Schedule (this is must)

Contracting Officer or POC Contact Info

Any attachments or referenced FAR clauses

🔎 3. Identify and Understand the Product/Service
Break down exactly what product or service is being requested.

Translate technical language or obscure item descriptions into standard commercial equivalents.

If applicable, explain any relevant specifications, standards (e.g., MIL-SPEC, ANSI, ISO), or compliance certifications required.

🛒 4. Source the Product or Service
Use sourcing logic to locate the exact or equivalent item(s) being requested.

Search GSA Advantage, DLA, SAM.gov vendor lists, and major commercial suppliers (e.g., Grainger, Fastenal, McKesson, Dell, CDW).

Include relevant part numbers, prices, and lead times.

🤝 5. Find Suitable and Compliant Suppliers
Recommend 2–3 legitimate, cost-effective suppliers who meet the government’s requirements.

Prefer vendors that are: U.S.-based, registered in SAM.gov, and classified as small business, 8(a), HUBZone, woman-owned, or veteran-owned (if the set-aside requires it).

Include links to supplier profiles or catalogs and reasons for each recommendation.

💰 6. Price Comparison and Cost Optimization
Compare all pricing options and identify the lowest possible cost that still meets quality and compliance standards.

Include total cost breakdown (unit cost, shipping, tax, etc.)

Suggest whether to quote direct, use a distributor, or purchase via an existing contract vehicle (e.g., GSA Schedule, NASA SEWP, DLA TLSP).

📋 7. Final Recommendations and Next Steps
Provide a summary of findings and a recommended sourcing path.

List next steps for me to take — e.g., gather compliance docs, contact supplier, submit capability statement, or prepare a quote.

If the opportunity is not viable, explain why and suggest alternatives.

📎 Format:
Organize all output using clear section headers, bullet points, tables (if needed), and logical flow. Keep technical language concise but accurate. If the solicitation is missing key info, ask clarifying questions.
    
    return your response in JSON with appropriate fields. Do not use any conversational cues just strictly provide the required information in this json way:
    <json field>: <related output text>
    """
}


def parse_raw_output(raw_output):
    try:
        # Step 1: Decode the outer string (remove escape characters)
        cleaned_str = bytes(raw_output, "utf-8").decode("unicode_escape")
        
        # Step 2: Strip any extra quotes that wrapped the object
        if cleaned_str.startswith('"') and cleaned_str.endswith('"'):
            cleaned_str = cleaned_str[1:-1]

        # Step 3: Now parse the inner JSON string
        result = json.loads(cleaned_str)
        return result
    except Exception as e:
        print("Failed to parse raw_output:", e)
        return None

DIRECT_UPLOAD_SUFFIXES = {
    "pdf": ".pdf",
    "docx": ".docx",
    "txt": ".txt",
    "md": ".md",
}

CONVERT_TO_PDF_TYPES = {"xlsx", "xls", "csv", "html"}


def _detect_file_type(url, content_type):
    """Detect file type from URL extension and Content-Type header."""
    from urllib.parse import urlparse
    ext = os.path.splitext(urlparse(url).path)[1].lower()
    
    type_map = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.xlsx': 'xlsx',
        '.xls': 'xls',
        '.html': 'html',
        '.htm': 'html',
        '.txt': 'txt',
        '.md': 'md',
        '.csv': 'csv',
    }
    if ext in type_map:
        return type_map[ext]
    
    ct = (content_type or '').lower()
    if 'pdf' in ct:
        return 'pdf'
    if 'wordprocessingml' in ct:
        return 'docx'
    if 'msword' in ct:
        return 'doc'
    if 'spreadsheetml' in ct:
        return 'xlsx'
    if 'ms-excel' in ct:
        return 'xls'
    if 'text/markdown' in ct:
        return 'md'
    if 'text/csv' in ct:
        return 'csv'
    if 'html' in ct:
        return 'html'
    if 'text/plain' in ct:
        return 'txt'
    
    return 'unknown'

def _sniff_file_type(file_content):
    """Best-effort file type sniffing for generic URLs/content-types."""
    if file_content[:5] == b'%PDF-':
        return 'pdf'

    # Office Open XML formats are ZIP containers.
    if file_content[:2] == b'PK':
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                names = zf.namelist()
                if any(name.startswith("word/") for name in names):
                    return 'docx'
                if any(name.startswith("xl/") for name in names):
                    return 'xlsx'
        except Exception:
            pass

    return 'unknown'


def _convert_to_pdf(file_content, file_type):
    """Convert non-PDF file content to PDF bytes. Returns PDF bytes or None."""
    try:
        if file_type == 'docx':
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    paragraphs.append(f"<p>{escaped}</p>")
            
            for table in doc.tables:
                paragraphs.append("<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse;width:100%'>")
                for row in table.rows:
                    paragraphs.append("<tr>")
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        paragraphs.append(f"<td>{cell_text}</td>")
                    paragraphs.append("</tr>")
                paragraphs.append("</table>")
            
            html = f"<html><body style='font-family:Arial,sans-serif;font-size:11pt;line-height:1.5'>{''.join(paragraphs)}</body></html>"
        
        elif file_type in ('xlsx', 'xls', 'csv'):
            if file_type == 'csv':
                import csv as csv_module
                reader = csv_module.reader(io.StringIO(file_content.decode('utf-8', errors='replace')))
                rows = list(reader)
                table_html = "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse;width:100%'>"
                for row in rows:
                    table_html += "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>"
                table_html += "</table>"
                html = f"<html><body style='font-family:Arial,sans-serif;font-size:10pt'>{table_html}</body></html>"
            else:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
                sheets_html = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheets_html.append(f"<h2>{sheet_name}</h2>")
                    sheets_html.append("<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse;width:100%'>")
                    for row in ws.iter_rows(values_only=True):
                        sheets_html.append("<tr>")
                        for cell in row:
                            val = str(cell) if cell is not None else ""
                            val = val.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            sheets_html.append(f"<td>{val}</td>")
                        sheets_html.append("</tr>")
                    sheets_html.append("</table>")
                wb.close()
                html = f"<html><body style='font-family:Arial,sans-serif;font-size:10pt'>{''.join(sheets_html)}</body></html>"
        
        elif file_type == 'html':
            html = file_content.decode('utf-8', errors='replace')
        
        elif file_type == 'txt':
            text = file_content.decode('utf-8', errors='replace')
            escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            html = f"<html><body style='font-family:monospace;font-size:10pt;white-space:pre-wrap'>{escaped}</body></html>"
        
        else:
            return None
        
        from weasyprint import HTML
        pdf_bytes = HTML(string=html).write_pdf()
        return pdf_bytes
    
    except Exception as e:
        print(f"Error converting {file_type} to PDF: {e}")
        return None


def download_and_upload_files(urls):
    uploaded_files = []
    stats = {
        "requested": len(urls),
        "uploaded": 0,
        "skipped": 0,
        "skipped_details": [],
    }
    for url in urls:
        tmp_path = None
        try:
            response = requests.get(url, timeout=60)
            response.raise_for_status()
            
            content_type = response.headers.get('Content-Type', '')
            file_type = _detect_file_type(url, content_type)
            sniffed_type = _sniff_file_type(response.content)
            if sniffed_type != 'unknown' and file_type in {'unknown', 'doc', 'xls'}:
                file_type = sniffed_type
            print(f"  File type detected: {file_type} for {url[:80]}...", flush=True)
            
            if file_type in DIRECT_UPLOAD_SUFFIXES:
                file_bytes = response.content
                upload_suffix = DIRECT_UPLOAD_SUFFIXES[file_type]
                print(f"  Uploading original {file_type} file", flush=True)
            elif file_type == 'doc':
                reason = "Legacy .doc format (unsupported)"
                stats["skipped"] += 1
                stats["skipped_details"].append({"url": url[:100], "reason": reason})
                print(f"  Skipping: {reason} for {url[:80]}...", flush=True)
                continue
            elif file_type in CONVERT_TO_PDF_TYPES:
                file_bytes = _convert_to_pdf(response.content, file_type)
                if not file_bytes:
                    reason = f"Failed to convert {file_type} to PDF"
                    stats["skipped"] += 1
                    stats["skipped_details"].append({"url": url[:100], "reason": reason})
                    print(f"  Skipping: {reason} for {url[:80]}...", flush=True)
                    continue
                upload_suffix = ".pdf"
                print(f"  Converted {file_type} to PDF ({len(file_bytes)} bytes)", flush=True)
            else:
                reason = f"Unsupported file type: {file_type}"
                stats["skipped"] += 1
                stats["skipped_details"].append({"url": url[:100], "reason": reason})
                print(f"  Skipping: {reason} for {url[:80]}...", flush=True)
                continue

            with tempfile.NamedTemporaryFile(delete=False, suffix=upload_suffix) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as f:
                upload = openai_client.files.create(
                    file=f,
                    purpose="user_data"
                )

            uploaded_files.append({
                "type": "input_file",
                "file_id": upload.id
            })
            stats["uploaded"] += 1

        except Exception as e:
            reason = str(e)
            stats["skipped"] += 1
            stats["skipped_details"].append({"url": url[:100], "reason": reason})
            print(f"Error downloading or uploading file from {url}: {e}", flush=True)
        finally:
            if tmp_path:
                try:
                    os.unlink(tmp_path)
                except Exception as del_err:
                    print(f"Warning: Could not delete temp file {tmp_path}: {del_err}", flush=True)
    return uploaded_files, stats


# Prompt for extracting key info from a single document
EXTRACTION_PROMPT = """
Extract information from this document and return it ONLY as a valid JSON object with exactly these 9 top-level keys.
Only extract what is explicitly present in this document. Use null for any field not found. Do not add any fields not listed below.

{
  "solicitation_metadata": {
    "solicitation_number": null,
    "amendment_numbers": [],
    "agency_name": null,
    "contracting_office": null,
    "naics_code": null,
    "psc_code": null,
    "set_aside_type": null,
    "contract_type": null,
    "response_deadline_date": null,
    "response_deadline_time": null,
    "time_zone": null,
    "posted_date": null,
    "period_of_performance": null,
    "place_of_performance": null,
    "unique_entity_id_required": null,
    "sam_registration_required": null
  },
  "submission_requirements": {
    "submission_method": null,
    "submission_email_address": null,
    "subject_line_format_requirement": null,
    "file_naming_requirements": null,
    "page_limits": null,
    "font_formatting_requirements": null,
    "required_attachments_list": [],
    "required_forms": [],
    "amendment_acknowledgment_required": null,
    "deadline_time_zone": null,
    "late_submission_policy": null,
    "questions_due_date": null,
    "questions_submission_method": null
  },
  "pricing_clin_information": {
    "clin_structure": [],
    "option_years": null,
    "ceiling_value": null,
    "funding_type": null,
    "pricing_spreadsheet_provided": null,
    "required_price_breakdown_format": null,
    "fob_terms": null
  },
  "technical_requirements": {
    "scope_of_work_summary": null,
    "detailed_technical_specifications": [],
    "brand_name_requirement": null,
    "brand_name_or_equal": null,
    "salient_characteristics": [],
    "installation_required": null,
    "warranty_requirements": null,
    "required_certifications": [],
    "packaging_requirements": null,
    "delivery_timeline": null,
    "security_clearance_required": null,
    "key_personnel_requirement": null,
    "past_performance_requirement": null,
    "minimum_years_of_experience": null
  },
  "evaluation_criteria": {
    "evaluation_method": null,
    "evaluation_factors": [],
    "price_weight_percent": null,
    "technical_weight_percent": null,
    "past_performance_weight_percent": null,
    "basis_of_award": null,
    "price_realism_analysis_required": null
  },
  "compliance_regulatory": {
    "far_clauses": [],
    "dfars_clauses": [],
    "buy_american_required": null,
    "trade_agreements_act_required": null,
    "itar_ear_restrictions": null,
    "cmmc_nist_cybersecurity_requirements": null,
    "insurance_requirements": null,
    "bonding_requirements": null,
    "wage_determination_included": null
  },
  "delivery_information": {
    "ship_to_address": null,
    "delivery_days_aro": null,
    "installation_location": null,
    "special_handling_requirements": null,
    "government_furnished_equipment": null
  },
  "amendments_attachments": {
    "number_of_amendments": null,
    "amendment_summaries": [],
    "amendment_acknowledgment_required": null,
    "list_of_attachments": [],
    "pricing_template_attachment": null,
    "drawings_technical_files_included": null,
    "qa_document_included": null
  },
  "disqualification_risk_factors": {
    "mandatory_site_visit": null,
    "exact_model_required": null,
    "extremely_short_delivery_window": null,
    "strict_formatting_requirements": null,
    "bonding_threshold": null,
    "insurance_threshold": null,
    "security_clearance_required": null,
    "multiple_mandatory_forms": null,
    "registration_in_additional_systems_required": null
  }
}

Return ONLY the filled-in JSON. No explanation, no markdown, no extra text.
"""

# Prompt for combining extracted data into final merged summary
FINAL_SUMMARY_PROMPT = """
You are merging multiple JSON extractions from different solicitation documents into one final JSON.
All input JSONs share exactly this schema with 9 top-level keys:
solicitation_metadata, submission_requirements, pricing_clin_information, technical_requirements,
evaluation_criteria, compliance_regulatory, delivery_information, amendments_attachments, disqualification_risk_factors.

Merge rules:
- For scalar fields (strings, Yes/No, numbers): use the first non-null value found across all documents.
- For list fields (arrays): combine all values from all documents into one deduplicated list.
- For narrative/summary fields (scope_of_work_summary, detailed_technical_specifications, etc.): synthesize all non-null values into one clear, comprehensive value.
- If a field is null in all documents, keep it null.
- Do NOT add any fields not in the schema. Do NOT remove any fields.
- Return ONLY valid JSON. No explanation, no markdown, no extra text.

Input data (list of per-document JSONs):
"""


def _is_rate_limit_error(e):
    """Check if exception is a 429 rate limit error — these should be retried"""
    err_str = str(e).lower()
    return "429" in str(e) or "rate_limit" in err_str or "rate limit" in err_str


def _is_retryable_error(e):
    """For final summary only: retry on both 429 and 500"""
    err_str = str(e).lower()
    return (
        "429" in str(e) or "rate_limit" in err_str or "rate limit" in err_str
        or "500" in str(e) or "server_error" in err_str
    )


def analyze_single_document(file_input):
    """Analyze a single document and extract key information"""
    max_retries = 4
    backoff_secs = [2, 5, 10, 20]

    for attempt in range(max_retries):
        try:
            response = openai_client.responses.create(
                model="gpt-4o-mini",  # Use mini for individual doc extraction (cheaper & faster)
                input=[
                    {
                        "role": "user",
                        "content": [
                            file_input,
                            {"type": "input_text", "text": EXTRACTION_PROMPT}
                        ]
                    }
                ]
            )
            raw_output = response.output_text
            # Try to parse JSON
            try:
                # Clean up the response - remove markdown code blocks if present
                cleaned = raw_output.strip()
                if cleaned.startswith("```json"):
                    cleaned = cleaned[7:]
                if cleaned.startswith("```"):
                    cleaned = cleaned[3:]
                if cleaned.endswith("```"):
                    cleaned = cleaned[:-3]
                return json.loads(cleaned.strip())
            except Exception:
                return {"raw_text": raw_output}
        except Exception as e:
            if _is_rate_limit_error(e) and attempt < max_retries - 1:
                wait = backoff_secs[attempt]
                print(f"  Rate limit hit, retrying in {wait}s (attempt {attempt + 1}/{max_retries})...", flush=True)
                time.sleep(wait)
            else:
                if not _is_rate_limit_error(e):
                    print(f"  Server error on document, skipping: {e}", flush=True)
                else:
                    print(f"  Rate limit exhausted, skipping document: {e}", flush=True)
                return {"error": str(e)}


def create_final_summary(extracted_data_list):
    """Combine extracted data from all documents into final summary"""
    raw_output = None
    max_retries = 4
    backoff_secs = [2, 5, 10, 20]

    for attempt in range(max_retries):
        try:
            combined_data = json.dumps(extracted_data_list, indent=2)

            response = openai_client.chat.completions.create(
                model="gpt-4o",  # Use full model for final synthesis
                messages=[
                    {
                        "role": "user",
                        "content": FINAL_SUMMARY_PROMPT + combined_data
                    }
                ],
                max_tokens=16000,
                temperature=0.1
            )

            raw_output = response.choices[0].message.content
            # Clean and parse
            cleaned = raw_output.strip()
            if cleaned.startswith("```json"):
                cleaned = cleaned[7:]
            if cleaned.startswith("```"):
                cleaned = cleaned[3:]
            if cleaned.endswith("```"):
                cleaned = cleaned[:-3]
            return json.loads(cleaned.strip())
        except json.JSONDecodeError as e:
            preview = (raw_output[:500] + "...") if raw_output and len(raw_output) > 500 else raw_output
            print(f"[create_final_summary] JSONDecodeError: {e}", flush=True)
            print(f"[create_final_summary] Raw output preview: {preview}", flush=True)
            return {"error": f"Invalid JSON from model: {e}", "raw_output_preview": preview}
        except Exception as e:
            if _is_retryable_error(e) and attempt < max_retries - 1:
                wait = backoff_secs[attempt]
                print(f"[create_final_summary] Retryable error (429/500), retrying in {wait}s (attempt {attempt + 1}/{max_retries})...", flush=True)
                time.sleep(wait)
            else:
                import traceback
                print(f"[create_final_summary] Exception: {e}", flush=True)
                print(f"[create_final_summary] Traceback: {traceback.format_exc()}", flush=True)
                return {"error": str(e)}


@app.route("/analyze-solicitations", methods=["POST"])
def analyze_solicitations():
    data = request.get_json()
    urls = data.get("urls")

    if not urls or not isinstance(urls, list):
        return jsonify({"error": "Missing or invalid 'urls' array."}), 400

    print(f"[analyze-solicitations] Documents requested: {len(urls)}", flush=True)

    # Upload files to OpenAI
    file_inputs, upload_stats = download_and_upload_files(urls)

    # Log upload stats
    print(f"[analyze-solicitations] Upload phase: requested={upload_stats['requested']}, "
          f"uploaded={upload_stats['uploaded']}, skipped={upload_stats['skipped']}", flush=True)
    if upload_stats["skipped_details"]:
        for d in upload_stats["skipped_details"]:
            print(f"  - Skipped: {d['reason']} | {d['url']}...", flush=True)

    if not file_inputs:
        return jsonify({
            "error": "Failed to upload any files.",
            "processing_stats": {
                "documents_requested": upload_stats["requested"],
                "documents_uploaded": 0,
                "documents_skipped": upload_stats["skipped"],
                "documents_analyzed": 0,
                "documents_in_summary": 0,
                "skipped_details": upload_stats["skipped_details"],
            }
        }), 500

    # Step 1: Process each document separately to extract key info
    print(f"[analyze-solicitations] Step 1: Extracting key information from {len(file_inputs)} document(s)...", flush=True)
    extracted_data = []
    for i, file_input in enumerate(file_inputs):
        if i > 0:
            time.sleep(3)  # Delay between documents to avoid rate limits
        print(f"  Processing document {i+1}/{len(file_inputs)}...", flush=True)
        doc_data = analyze_single_document(file_input)
        doc_data["document_index"] = i + 1
        extracted_data.append(doc_data)

    # Filter out documents that failed to analyze
    successful_data = [d for d in extracted_data if "error" not in d]
    failed_count = len(extracted_data) - len(successful_data)

    print(f"[analyze-solicitations] Extraction phase: analyzed={len(successful_data)}/{len(file_inputs)}, "
          f"failed={failed_count}", flush=True)
    if failed_count > 0:
        for i, d in enumerate(extracted_data):
            if "error" in d:
                print(f"  - Document {i+1} failed: {d.get('error', 'unknown')}", flush=True)

    if not successful_data:
        return jsonify({
            "error": "All documents failed to analyze.",
            "processing_stats": {
                "documents_requested": upload_stats["requested"],
                "documents_uploaded": upload_stats["uploaded"],
                "documents_skipped": upload_stats["skipped"],
                "documents_analyzed": 0,
                "documents_in_summary": 0,
                "skipped_details": upload_stats["skipped_details"],
            }
        }), 500

    # Step 2: Combine all extracted data into final summary
    print(f"[analyze-solicitations] Step 2: Creating final summary from {len(successful_data)} document(s)...", flush=True)
    try:
        final_summary = create_final_summary(successful_data)
        processing_stats = {
            "documents_requested": upload_stats["requested"],
            "documents_uploaded": upload_stats["uploaded"],
            "documents_skipped": upload_stats["skipped"],
            "documents_analyzed": len(successful_data),
            "documents_in_summary": len(successful_data),
            "skipped_details": upload_stats["skipped_details"],
        }
        final_summary["processing_stats"] = processing_stats
        print(f"[analyze-solicitations] Done. Summary generated from {len(successful_data)}/{upload_stats['requested']} "
              f"requested documents.", flush=True)
        return jsonify(final_summary)
    except Exception as e:
        return jsonify({
            "error": str(e),
            "extracted_data": extracted_data,
            "processing_stats": {
                "documents_requested": upload_stats["requested"],
                "documents_uploaded": upload_stats["uploaded"],
                "documents_skipped": upload_stats["skipped"],
                "documents_analyzed": len(successful_data),
                "documents_in_summary": 0,
                "skipped_details": upload_stats["skipped_details"],
            }
        }), 500


@app.route("/message-chat", methods=["POST"])
def message_chat():
    data = request.get_json()
    
    summary = data.get("summary")
    chat_history = data.get("chatHistory", [])
    user_message = data.get("userMessage")
    
    if not user_message:
        return jsonify({"error": "Missing 'userMessage' field."}), 400
    
    # Build conversation context
    system_message = f"""You are a helpful government contracting assistant. You have access to the following contract summary information:

{json.dumps(summary, indent=2) if summary else "No summary available."}

Use this context to answer questions about the contract. Be helpful, accurate, and concise. If you don't know something or it's not in the provided context, say so."""

    # Build messages array for OpenAI
    messages = [{"role": "system", "content": system_message}]
    
    # Add chat history
    for msg in chat_history:
        role = "assistant" if msg.get("role") == "agent" else "user"
        messages.append({"role": role, "content": msg.get("content", "")})
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=1000,
            temperature=0.7
        )
        
        assistant_message = response.choices[0].message.content
        return jsonify({"message": assistant_message})
    
    except Exception as e:
        print(f"Error in message-chat: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    with app.app_context():
        # Import models first
        from quote import NegotiationSession, Supplier, Message
        db.create_all()
        
        # Initialize background jobs
        init_background_jobs(db, Supplier, Message, NegotiationSession, app.app_context)
    
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 9000)), debug=True, use_reloader=False)

